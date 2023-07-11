import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH    #对齐方式
from docx.oxml.ns import qn       # 字体设置函数
from docx.shared import Inches,Pt   # 导入缩进单位

i=0

text=["祝你永远年轻，永远在路上，永远热泪盈眶",
     "和所有的烦恼说拜拜，和所有的快乐说嗨嗨",
     "愿你一切尽意，百事从欢",
	 "愿你心怀坦荡，豁达开朗，有酒有诗有远方",
	 "愿你想要的都得到，得到的都美好"
	]

picture=["a.jpg","b.jpg","c.jpg","d.jpg","e.jpg"]

def makeDoc(name,gender):
	global i
	doc = docx.Document()
	para=doc.add_heading(name,0)    #标题一的格式
	if gender=='男':
		para.add_run("先生：")
	else:
		para.add_run("女士：")
	para.runs[0].font.size=Pt(18)      #名字，字号18磅
	para.runs[1].font.size=Pt(18)      #称呼，字号18磅
	doc.add_paragraph(text[i%5])       #i%5表示，每五个人一组，一组中的文案是不相同的
	p = doc.paragraphs[1]
	p.paragraph_format.left_indent = Inches(0.5)  #第1段左缩进0.5英寸
	doc.styles['Normal'].font.name = u'宋体'     #设置字体
	doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
	doc.add_picture(picture[i%5], width=Inches(4))            #添加图片
	doc.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER   #图片居中
	i+=1
	doc.add_paragraph("您的好友：杨宇婵").alignment = WD_ALIGN_PARAGRAPH.RIGHT    #署名靠右对齐
	doc.save(f'给{name}的贺卡.docx')

f = open("name.txt",'r',encoding='utf-8')   #打开文件
content = f.readlines()
for line in content:
	line=line.strip('\n')   #去掉每行结束的换行符\n
	splitstr=line.split()   #以空格将字符串切片
	name=splitstr[0]
	gender=splitstr[1]
	makeDoc(name,gender)    #制作贺卡的函数

